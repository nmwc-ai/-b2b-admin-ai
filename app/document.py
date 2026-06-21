"""견적서/계약서 DOCX 생성 — 서버리스용(메모리 BytesIO 반환, 파일 저장 없음).

기존 document.py를 이식하되 OUTPUT_DIR 파일 저장 대신 bytes를 반환해 다운로드로 스트리밍.
"""

import io
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app import settings


def _antiegg_info() -> dict:
    return {
        'ceo':    settings.get('ANTIEGG_CEO', 'ANTIEGG 대표'),
        'biz_no': settings.get('ANTIEGG_BIZ_NO', '000-00-00000'),
        'phone':  settings.get('ANTIEGG_PHONE'),
        'email':  settings.get('ANTIEGG_EMAIL') or settings.get('DIRECTOR_EMAIL'),
        'addr':   settings.get('ANTIEGG_ADDR'),
    }


def _set_font(run, size=11, bold=False, color=None):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _cell_text(cell, text, size=10, bold=False, align='left', color=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                      'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(str(text))
    _set_font(run, size=size, bold=bold, color=color)


def _shade_cell(cell, hex_color='1F3864'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])


def _parse_number(value) -> int:
    cleaned = ''.join(c for c in str(value or '') if c.isdigit())
    return int(cleaned) if cleaned else 0


def _fmt_money(n: int) -> str:
    return f'{n:,}원'


def _save_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── 견적서 ───────────────────────────────────────────────────────────────────
def generate_quote(deal: dict) -> bytes:
    ant = _antiegg_info()
    unit_price = _parse_number(deal.get('cond_unit_price'))
    quantity = _parse_number(deal.get('cond_quantity')) or 1
    subtotal = unit_price * quantity
    vat = int(subtotal * 0.1)
    total = subtotal + vat
    today = datetime.now()
    quote_date = today.strftime('%Y년 %m월 %d일')
    valid_str = (today + timedelta(days=30)).strftime('%Y년 %m월 %d일')

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.0)
        s.left_margin = s.right_margin = Cm(2.5)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(tp.add_run('견  적  서'), size=22, bold=True, color=(31, 56, 100))
    doc.add_paragraph()

    info = doc.add_table(rows=4, cols=4)
    info.style = 'Table Grid'
    _set_col_widths(info, [2.5, 5.5, 2.5, 5.5])

    def info_row(r, l1, v1, l2, v2):
        c = info.rows[r].cells
        _cell_text(c[0], l1, bold=True, align='center', color=(255, 255, 255)); _shade_cell(c[0])
        _cell_text(c[1], v1)
        _cell_text(c[2], l2, bold=True, align='center', color=(255, 255, 255)); _shade_cell(c[2])
        _cell_text(c[3], v2)

    info_row(0, '수 신', deal.get('company', '') or '', '견적번호', deal['deal_id'])
    info_row(1, '담당자', deal.get('contact_name', '') or '', '견적일', quote_date)
    info_row(2, '이메일', deal.get('email', '') or '', '유효기간', valid_str)
    info_row(3, '연락처', deal.get('contact_phone', '') or '', '발신', ant['ceo'])
    doc.add_paragraph()

    doc.add_heading('서비스 명세', level=2)
    it = doc.add_table(rows=2, cols=5)
    it.style = 'Table Grid'
    _set_col_widths(it, [5.0, 5.0, 2.5, 2.0, 3.0])
    for i, h in enumerate(['서비스명', '서비스 설명', '단가', '수량', '공급가액']):
        _cell_text(it.rows[0].cells[i], h, bold=True, align='center', color=(255, 255, 255))
        _shade_cell(it.rows[0].cells[i])
    r1 = it.rows[1].cells
    _cell_text(r1[0], deal.get('cond_service_name', '') or '')
    _cell_text(r1[1], deal.get('cond_service_desc', '') or '')
    _cell_text(r1[2], _fmt_money(unit_price), align='right')
    _cell_text(r1[3], str(quantity), align='center')
    _cell_text(r1[4], _fmt_money(subtotal), align='right')
    doc.add_paragraph()

    tt = doc.add_table(rows=3, cols=2)
    tt.style = 'Table Grid'
    _set_col_widths(tt, [13.5, 4.0])
    for i, (label, value) in enumerate([('공급가액', _fmt_money(subtotal)),
                                        ('부가세 (10%)', _fmt_money(vat)), ('합  계', _fmt_money(total))]):
        c = tt.rows[i].cells
        is_total = (i == 2)
        col = (255, 255, 255) if is_total else None
        _cell_text(c[0], label, bold=is_total, align='right', color=col)
        _cell_text(c[1], value, bold=is_total, align='right', color=col)
        if is_total:
            _shade_cell(c[0]); _shade_cell(c[1])
    doc.add_paragraph()

    conds = [('결제 조건', deal.get('cond_payment_terms', '')), ('납품 범위', deal.get('cond_delivery_scope', '')),
             ('특이사항', deal.get('cond_notes', ''))]
    if any(v for _, v in conds):
        doc.add_heading('계약 조건', level=2)
        ct = doc.add_table(rows=len(conds), cols=2)
        ct.style = 'Table Grid'
        _set_col_widths(ct, [3.5, 14.0])
        for i, (label, value) in enumerate(conds):
            c = ct.rows[i].cells
            _cell_text(c[0], label, bold=True, align='center', color=(255, 255, 255)); _shade_cell(c[0], '2F5496')
            _cell_text(c[1], value or '—')
        doc.add_paragraph()

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lines = ['ANTIEGG']
    if ant['addr']:   lines.append(ant['addr'])
    if ant['biz_no']: lines.append(f"사업자등록번호: {ant['biz_no']}")
    if ant['phone']:  lines.append(f"Tel: {ant['phone']}")
    if ant['email']:  lines.append(f"Email: {ant['email']}")
    _set_font(sp.add_run('\n'.join(lines)), size=9, color=(89, 89, 89))

    return _save_bytes(doc)


# ── 계약서 ───────────────────────────────────────────────────────────────────
def _article(doc, title, body):
    tp = doc.add_paragraph()
    _set_font(tp.add_run(title), size=11, bold=True, color=(31, 56, 100))
    bp = doc.add_paragraph()
    _set_font(bp.add_run(body), size=10)


def generate_contract(deal: dict) -> bytes:
    ant = _antiegg_info()
    unit_price = _parse_number(deal.get('cond_unit_price'))
    quantity = _parse_number(deal.get('cond_quantity')) or 1
    subtotal = unit_price * quantity
    vat = int(subtotal * 0.1)
    total = subtotal + vat
    today_str = datetime.now().strftime('%Y년 %m월 %d일')
    cstart = deal.get('cond_contract_start') or '___년 ___월 ___일'
    cend = deal.get('cond_contract_end') or '___년 ___월 ___일'
    company = deal.get('company', '') or ''
    client_ceo = deal.get('cond_company_ceo', '') or ''
    client_biz = deal.get('cond_company_biz_no', '') or ''

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.5)
        s.left_margin = s.right_margin = Cm(3.0)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(tp.add_run('용  역  계  약  서'), size=20, bold=True, color=(31, 56, 100))
    doc.add_paragraph()

    def party_block(rows_data):
        t = doc.add_table(rows=1 + len(rows_data), cols=4)
        t.style = 'Table Grid'
        _set_col_widths(t, [1.8, 5.2, 1.8, 5.2])
        h = t.rows[0]
        h.cells[0].merge(h.cells[1]); h.cells[2].merge(h.cells[3])
        _cell_text(h.cells[0], '갑 (발주자)', bold=True, align='center', color=(255, 255, 255)); _shade_cell(h.cells[0])
        _cell_text(h.cells[2], '을 (수급자)', bold=True, align='center', color=(255, 255, 255)); _shade_cell(h.cells[2])
        for i, (la, va, lb, vb) in enumerate(rows_data, start=1):
            c = t.rows[i].cells
            _cell_text(c[0], la, bold=True, align='center'); _cell_text(c[1], va)
            _cell_text(c[2], lb, bold=True, align='center'); _cell_text(c[3], vb)

    party_block([
        ('상    호', company, '상    호', 'ANTIEGG'),
        ('대 표 자', client_ceo, '대 표 자', ant['ceo']),
        ('사업자번호', client_biz, '사업자번호', ant['biz_no']),
        ('연  락  처', deal.get('contact_phone', '') or deal.get('email', '') or '',
         '연  락  처', ant['phone'] or ant['email'] or ''),
    ])
    doc.add_paragraph()

    _article(doc, '제1조 (목적)',
             f'본 계약은 {company}(이하 "갑")와 ANTIEGG(이하 "을") 간에 '
             f'{deal.get("cond_service_name", "서비스") or "서비스"} 제공에 관한 권리와 의무를 규정함을 목적으로 한다.')
    _article(doc, '제2조 (용역 범위)',
             f'1. 서비스명: {deal.get("cond_service_name", "") or ""}\n'
             f'2. 내용: {deal.get("cond_service_desc", "") or ""}\n'
             f'3. 납품 범위: {deal.get("cond_delivery_scope", "") or "별도 협의"}')
    _article(doc, '제3조 (계약 기간)', f'계약 기간은 {cstart}부터 {cend}까지로 한다.')
    _article(doc, '제4조 (계약 금액)',
             f'공급가액: {_fmt_money(subtotal)}\n부가가치세 (10%): {_fmt_money(vat)}\n합계 (VAT 포함): {_fmt_money(total)}')
    _article(doc, '제5조 (결제 조건)', deal.get('cond_payment_terms', '') or '계약 체결 후 별도 협의')
    n = 6
    if deal.get('cond_notes'):
        _article(doc, f'제{n}조 (특이사항)', deal['cond_notes']); n += 1
    _article(doc, f'제{n}조 (일반 조항)',
             '① 본 계약에 명시되지 않은 사항은 관련 법령 및 상관례에 따른다.\n'
             '② 본 계약과 관련하여 분쟁이 발생한 경우 갑, 을이 협의하여 해결한다.')
    doc.add_paragraph()

    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(dp.add_run(today_str), size=11, bold=True)
    doc.add_paragraph()

    st = doc.add_table(rows=4, cols=4)
    st.style = 'Table Grid'
    _set_col_widths(st, [1.8, 5.2, 1.8, 5.2])
    h = st.rows[0]
    h.cells[0].merge(h.cells[1]); h.cells[2].merge(h.cells[3])
    _cell_text(h.cells[0], '갑 (발주자)', bold=True, align='center', color=(255, 255, 255)); _shade_cell(h.cells[0])
    _cell_text(h.cells[2], '을 (수급자)', bold=True, align='center', color=(255, 255, 255)); _shade_cell(h.cells[2])
    for i, (la, va, lb, vb) in enumerate([
        ('상    호', company, '상    호', 'ANTIEGG'),
        ('대 표 자', f'{client_ceo}  (인)', '대 표 자', f"{ant['ceo']}  (인)"),
        ('사업자번호', client_biz, '사업자번호', ant['biz_no']),
    ], start=1):
        c = st.rows[i].cells
        _cell_text(c[0], la, bold=True, align='center'); _cell_text(c[1], va)
        _cell_text(c[2], lb, bold=True, align='center'); _cell_text(c[3], vb)

    return _save_bytes(doc)
